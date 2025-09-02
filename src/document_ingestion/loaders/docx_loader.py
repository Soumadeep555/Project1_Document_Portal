# src/document_ingestion/loaders/docx_loader.py
from pathlib import Path
from zipfile import ZipFile
from docx import Document
import pandas as pd
from .base import BaseLoader, LoadedDocument, ImageBlob
from .image_utils import save_image_bytes


class DocxLoader(BaseLoader):
    def load(self, path: str) -> LoadedDocument:
        p = Path(path)
        doc = Document(str(p))
        text = "\n".join([p_.text for p_ in doc.paragraphs if p_.text])

        # tables
        tables = []
        for t in doc.tables:
            rows = []
            for row in t.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            if rows:
                header = rows[0] if any(rows[0]) else [f"col_{i}" for i in range(len(rows[0]))]
                data = rows[1:] if any(rows[0]) else rows
                tables.append(pd.DataFrame(data, columns=header))

        # images (from the .docx zip package under word/media/)
        images = []
        try:
            with ZipFile(str(p), "r") as zf:
                for name in zf.namelist():
                    if name.startswith("word/media/"):
                        bytes_ = zf.read(name)
                        out = save_image_bytes(bytes_, f"{p.stem}_{Path(name).name}")
                        images.append(ImageBlob(path=out, source=str(p)))
        except Exception:
            pass

        return LoadedDocument(source=str(p), text=text, tables=tables, images=images)
